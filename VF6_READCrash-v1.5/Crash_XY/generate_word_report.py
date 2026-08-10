"""
****************************************************************************
**  气囊控制器 (ACM/ACU) 冲击测试波形自动化分析报告生成脚本              **
**  1. 使用相对路径，自动定位当前脚本所在的 XY 根目录                      **
**  2. 一次性扫描并汇总 XY 下所有子目录 (data15, data30...) 的 Excel 文件     **
**  3. 自动生成包含波形图(高15cm x 宽30cm)、±5%极值校验及 30cm 表格的报告     **
****************************************************************************
"""

import os
import re
import tempfile
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 支持 Matplotlib 中文显示
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def setup_document_landscape(doc):
    """
    设置 Word 文档页面布局：
    页面总宽 35cm，左右边距各 2.5cm，版面显示区域刚好为 30cm，与图表及表格 1:1 完美契合
    """
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(35)
    section.page_height = Cm(25)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def analyze_waveform_data(df):
    """
    【最高点与最低点独立校验算法】：
    对 最高点 (np.max) 与 最低点 (np.min) 分别计算：
    1. 均值；
    2. [均值 - |均值|*5%, 均值 + |均值|*5%] 的合格区间；
    3. 校验每个通道的测试值是否全部落入该区间。
    """
    cols = df.columns
    if len(cols) == 0:
        return None

    p1_list, t1_list = [], []

    # 1. 提取每条通道的最高点与最低点
    for col in cols:
        vals = pd.to_numeric(df[col], errors='coerce').dropna().values
        if len(vals) == 0:
            continue

        p1 = float(np.max(vals))  # 最高点 (峰值)
        t1 = float(np.min(vals))  # 最低点 (谷值)

        p1_list.append(p1)
        t1_list.append(t1)

    if not p1_list:
        return None

    # 辅助校验函数：计算均值、正负 5% 区间及判定超差通道
    def check_metric_range(vals_list):
        avg = float(np.mean(vals_list))
        tol = abs(avg) * 0.05  # 正负 5% 门限
        min_b, max_b = avg - tol, avg + tol
        abnormal = [i+1 for i, v in enumerate(vals_list) if not (min_b <= v <= max_b)]
        is_ok = (len(abnormal) == 0)
        return avg, min_b, max_b, is_ok, abnormal

    # 2. 对最高点与最低点计算均值、正负 5% 区间和判定
    p1_avg, p1_min, p1_max, p1_ok, p1_ab = check_metric_range(p1_list)
    t1_avg, t1_min, t1_max, t1_ok, t1_ab = check_metric_range(t1_list)

    # 2 项指标均通过才算整体正常
    is_all_normal = p1_ok and t1_ok

    return {
        "p1": {"list": p1_list, "avg": p1_avg, "min": p1_min, "max": p1_max, "ok": p1_ok, "abnormal": p1_ab},
        "t1": {"list": t1_list, "avg": t1_avg, "min": t1_min, "max": t1_max, "ok": t1_ok, "abnormal": t1_ab},
        "is_normal": is_all_normal,
        "series_count": len(p1_list)
    }

def generate_chart_image(df, title, temp_dir):
    """绘制高 15cm, 宽 30cm 的波形对比图"""
    fig_width = 30 / 2.54   # 宽度 30cm
    fig_height = 15 / 2.54  # 高度 15cm

    fig, ax = plt.subplots(figsize=(fig_width, fig_height), dpi=150)

    for col in df.columns:
        series_data = pd.to_numeric(df[col], errors='coerce').dropna()
        if len(series_data) > 0:
            ax.plot(series_data.index + 1, series_data.values, label=str(col)[:40], linewidth=1.5)

    ax.set_title(f"{title} - 脉冲响应对比图", fontsize=15, fontweight='bold', pad=12)
    ax.set_xlabel("采样点序号 (Points)", fontsize=11)
    ax.set_ylabel("加速度物理幅值 (g)", fontsize=11)
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.axhline(0, color='black', linewidth=0.8, linestyle='-')

    if len(df.columns) <= 10:
        ax.legend(loc='upper right', fontsize=8.5, frameon=True)
    else:
        ax.legend(loc='upper right', bbox_to_anchor=(1.12, 1.0), fontsize=7.5, frameon=True)

    plt.tight_layout()

    safe_title = re.sub(r'[\/:*?"<>|]', '_', title)
    img_path = os.path.join(temp_dir, f"chart_{safe_title}.png")
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return img_path

def add_summary_table(doc, stats):
    """
    在 Word 中插入最高点与最低点正负 5% 区间校验表格
    【自适应宽度设计】：总宽度精确设为 30.0 cm，与上方 30cm 图表左右完全对齐
    """
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中对齐
    table.style = 'Table Grid'

    # 各列宽度分配，总和 = 6.0 + 5.0 + 7.5 + 6.5 + 5.0 = 30.0 cm
    col_widths = [Cm(6.0), Cm(5.0), Cm(7.5), Cm(6.5), Cm(5.0)]

    headers = ["极值指标项", "实测均值 (g)", "正负 5% 合格区间范围", "超差通道", "单项评估"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        # 表头深蓝背景填充
        shading = parse_xml(r'<w:shd {} w:fill="1F497D"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

    def get_status_str(item):
        return "合格 (正常)" if item['ok'] else "超差 (异常)"

    def get_ab_str(item):
        return "无" if item['ok'] else f"通道 {item['abnormal']}"

    data_rows = [
        ("最高点 (峰值)", f"{stats['p1']['avg']:.2f}", f"{stats['p1']['min']:.3f} ～ {stats['p1']['max']:.3f}", get_ab_str(stats['p1']), get_status_str(stats['p1'])),
        ("最低点 (谷值)", f"{stats['t1']['avg']:.2f}", f"{stats['t1']['min']:.3f} ～ {stats['t1']['max']:.3f}", get_ab_str(stats['t1']), get_status_str(stats['t1'])),
        ("通道数量", f"{stats['series_count']} 条", "综合判定结论", "全在区间内" if stats['is_normal'] else "存在超差", "【正常】" if stats['is_normal'] else "【异常】")
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9.0)

    # 统一设置每一列的精确宽度，保证全表宽度正好为 30cm
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph()  # 下方空行间隔

def extract_data_num(path):
    """提取 data15, data30, data45 中的数字，用于自然升序排列"""
    match = re.search(r'data(\d+)', path, re.IGNORECASE)
    return int(match.group(1)) if match else 999

def generate_all_reports(xy_root_dir, output_docx_path):
    """主函数：自动递归扫描 XY 下所有 dataXX 目录并汇总生成一份 Word 报告"""
    if not os.path.exists(xy_root_dir):
        print(f"错误: 根目录不存在 -> {xy_root_dir}")
        return

    # 扫描 XY 根目录及其所有子目录下所有的 .xlsx 文件
    found_excel_files = []
    for root, dirs, files in os.walk(xy_root_dir):
        for file in files:
            if file.lower().endswith('.xlsx') and not file.startswith('~$'):
                full_path = os.path.join(root, file)
                found_excel_files.append(full_path)

    if not found_excel_files:
        print(f"在目录 {xy_root_dir} 及其子目录下未找到任何 .xlsx 文件！")
        return

    # 按照 data15, data30, data45... 自然数字顺序排序
    found_excel_files.sort(key=extract_data_num)

    print("="*60)
    print(f"【批量生成 Word 报告】在 XY 目录下共扫描到 {len(found_excel_files)} 个 Excel 数据包：")
    for f in found_excel_files:
        print(f"  └─ {f}")
    print("="*60 + "\n开始汇总计算并生成 Word 报告...\n")

    doc = Document()
    setup_document_landscape(doc)

    # 报告大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("气囊控制器 (ACM/ACU) 冲击测试波形全角度综合汇总分析报告")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    # 试验背景与规则说明
    desc_p = doc.add_paragraph()
    desc_p.add_run("试验环境与判定规则说明：\n").bold = True
    desc_p.add_run("1. 涵盖多角度（如 15°、30°、45°、60°、75°、90° 等）全套冲击测试数据。\n")
    desc_p.add_run("2. 槽位与 IP 映射：1号槽位(172.16.3.171)、2号槽位(172.16.3.172)、3号槽位(172.16.3.173)、4号槽位(172.16.3.174)。\n")
    desc_p.add_run("3. 判定规则：分别提取最高点与最低点计算均值及 [均值×95% ~ 平均值×105%] 合格区间，若所有测试值全落在各自区间内判定为【正常】，否则判定为【异常】。\n")
    doc.add_paragraph()

    temp_dir = tempfile.mkdtemp()

    for excel_path in found_excel_files:
        rel_folder = os.path.basename(os.path.dirname(excel_path))
        excel_file = os.path.basename(excel_path)

        print(f"\n正在处理 [{rel_folder}] 中的数据包: {excel_file} ...")

        # =====================================================================
        # 一级目录：以 [子文件夹名 - Excel文件名] 命名
        # =====================================================================
        h1 = doc.add_heading(f"测试组 [{rel_folder}] : {excel_file}", level=1)
        h1.runs[0].font.color.rgb = RGBColor(31, 73, 125)
        h1.runs[0].font.size = Pt(16)

        try:
            xl = pd.ExcelFile(excel_path)
        except Exception as e:
            print(f"  └─ 读取 Excel 失败 [{excel_file}]: {str(e)}")
            continue

        for sheet_name in xl.sheet_names:
            print(f"  └─ 处理二级目录 Sheet: {sheet_name}")

            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            if df.empty or len(df.columns) == 0:
                continue

            # =================================================================
            # 二级目录：以 Sheet 名称命名
            # =================================================================
            h2 = doc.add_heading(f"波形对比分组: {sheet_name}", level=2)
            h2.runs[0].font.color.rgb = RGBColor(68, 114, 196)
            h2.runs[0].font.size = Pt(13)

            # 1. 最高点与最低点正负 5% 门限校验
            stats = analyze_waveform_data(df)
            if not stats:
                continue

            # 2. 生成波形图并插入 Word (高度: 15cm, 宽度: 30cm)
            chart_img_path = generate_chart_image(df, f"{rel_folder}_{sheet_name}", temp_dir)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(chart_img_path, height=Cm(15), width=Cm(30))

            # 3. 100% 完整列出所有通道的测试值，无任何省略
            p1_str = ", ".join([f"{v:.1f}" for v in stats['p1']['list']])
            t1_str = ", ".join([f"{v:.1f}" for v in stats['t1']['list']])

            status_tag = "【正常】" if stats['is_normal'] else "【异常】"
            eval_desc = "最高点与最低点实测值均 100% 落在各自 ±5% 合格区间内，波动平稳，一致性良好。" if stats['is_normal'] else "部分指标测试值超出 ±5% 门限区间，存在偏差，请关注对应超差通道。"

            narrative = (
                f"【最高点与最低点正负 5% 区间判定】\n"
                f"本组（{sheet_name}）共包含 {stats['series_count']} 条对比波形。针对最高点与最低点的判定如下：\n"
                f"1. 最高点 (峰值)：测试值 [{p1_str}]，计算均值为 {stats['p1']['avg']:.2f} g，正负 5% 合格区间为 [{stats['p1']['min']:.3f} ～ {stats['p1']['max']:.3f}] g → 单项判定: {'【合格】' if stats['p1']['ok'] else '【超差】'}；\n"
                f"2. 最低点 (谷值)：测试值 [{t1_str}]，计算均值为 {stats['t1']['avg']:.2f} g，正负 5% 合格区间为 [{stats['t1']['min']:.3f} ～ {stats['t1']['max']:.3f}] g → 单项判定: {'【合格】' if stats['t1']['ok'] else '【超差】'}。\n"
                f"综合判定结论：{status_tag}。{eval_desc}"
            )

            analysis_p = doc.add_paragraph()
            analysis_p.paragraph_format.line_spacing = 1.25
            run_narrative = analysis_p.add_run(narrative)
            run_narrative.font.size = Pt(10.0)

            # 4. 插入宽度 30cm 的总结表格
            add_summary_table(doc, stats)

    # 5. 保存汇总 Word 报告到 XY 根目录下
    try:
        doc.save(output_docx_path)
        print("\n" + "="*60)
        print(f"🎉 综合 Word 报告已成功生成至 XY 根目录：\n{output_docx_path}")
        print("="*60)
    except PermissionError:
        import datetime
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        backup_docx = os.path.join(xy_root_dir, f"冲击测试波形全角度汇总分析报告_{timestamp}.docx")
        print(f"\n[⚠️ 警告] 目标 Word 文件正在被打开占用，已自动保存至备用文件：\n{backup_docx}")
        doc.save(backup_docx)

if __name__ == "__main__":
    # 【相对路径】：自动获取当前脚本 generate_word_report.py 所在的 XY 根目录
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

    # 汇总 Word 报告存放在 XY 根目录下
    OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "冲击测试波形全角度汇总分析报告.docx")

    generate_all_reports(SCRIPT_DIR, OUTPUT_DOCX)
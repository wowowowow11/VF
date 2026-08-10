"""
****************************************************************************
**  代码名称：ACM_ImpactTest_AutoProcessor.py                            **
**  中文名称：气囊控制器 (ACM/ACU) 冲击测试全流程自动化处理系统            **
**                                                                        **
**  总体功能：                                                            **
**  1. 【TXT -> JSON】：自动解包所有 txt 为单通道 500 点波形 JSON            **
**  2. 【JSON -> Excel】：自动合并各子目录数据，生成多 Sheet + 15x30cm 折线图 **
**  3. 【Excel -> Word】：自动计算正负5%极值门限，生成含全数据分析的 Word 报告  **
****************************************************************************
"""

import os
import re
import json
import datetime
import tempfile
import openpyxl
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.chart import LineChart, Reference

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 支持 Matplotlib 中文显示
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


# =============================================================================
# 第一部分：TXT 文本数据 -> 500 点波形 JSON 解析模块
# =============================================================================

def parse_single_channel_hex(hex_str):
    """将整行十六进制数据按两两一组 (小端序 int16) 全量解包为 500 个采样点"""
    raw_bytes = [int(b, 16) for b in hex_str.split()]
    acc_values = []
    for i in range(0, len(raw_bytes) - 1, 2):
        raw = bytes([raw_bytes[i + 1], raw_bytes[i]])
        val = int.from_bytes(raw, byteorder='big', signed=True)
        acc_values.append(val)
    return acc_values


def sanitize_filename(name):
    """替换 Windows 文件名非法字符"""
    return re.sub(r'[\/:*?"<>|]', '-', name)


def process_txt_folder(target_dir):
    """处理单个文件夹下的所有 TXT 文件并生成对应的 JSON"""
    files = [f for f in os.listdir(target_dir) if f.lower().endswith('.txt')]
    if not files:
        return 0

    print(f"  ├─ [1/3 TXT解析] 找到 {len(files)} 个 TXT 文件，开始解包...")
    json_count = 0

    for file_name in files:
        txt_file_path = os.path.join(target_dir, file_name)
        txt_filename_base = os.path.splitext(file_name)[0]

        with open(txt_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line = line.strip()
            if not line or ",Data:" not in line:
                continue

            try:
                meta_part, data_part = line.split(",Data:", 1)
                meta_info = meta_part.strip()
                hex_str = data_part.strip()

                acc_values = parse_single_channel_hex(hex_str)

                if "Crash X" in meta_info:
                    parsed_result = {"000_x_acc_variable": acc_values}
                elif "Crash Y" in meta_info:
                    parsed_result = {"001_y_acc_variable": acc_values}
                else:
                    parsed_result = {"acc_variable": acc_values}

                clean_meta = sanitize_filename(meta_info)
                json_filename = f"{txt_filename_base}_{clean_meta}.json"
                json_save_path = os.path.join(target_dir, json_filename)

                with open(json_save_path, 'w', encoding='utf-8') as jf:
                    json.dump(parsed_result, jf, indent=4, ensure_ascii=False)

                json_count += 1
            except Exception as e:
                print(f"  │  └─ 解析失败 [{meta_info}]: {str(e)}")

    print(f"  │  └─ 解析完成，生成 {json_count} 个 JSON 文件。")
    return json_count


# =============================================================================
# 第二部分：JSON 文件 -> Excel 汇总与绘图模块
# =============================================================================

def parse_meta_from_json_filename(filename):
    """从 JSON 文件名解析元数据"""
    pattern = r"^(.*?)_\[(.*?)\]\s*,\s*ACU_SN[:\-]?\s*(.*?)\s*,\s*(Crash\s*[\w\s]+)\.json"
    match = re.search(pattern, filename, re.IGNORECASE)
    if match:
        return match.group(1).strip(), match.group(2).strip(), match.group(3).strip(), match.group(4).strip()
    return filename, "", "", "Crash X"


def make_smart_sheet_title(base_file, acu_sn, crash_type):
    """生成符合 Excel 31 字符上限规则的 Sheet 名称"""
    clean_sn = re.sub(r'\s+', '', acu_sn)
    clean_crash = crash_type.strip()
    ip_match = re.search(r"(\d+\.\d+\.\d+\.\d+|\d+)", base_file)
    short_file = ip_match.group(1) if ip_match else base_file[:10]

    candidate = f"{short_file}_{clean_sn}_{clean_crash}"
    candidate = re.sub(r'[\/:*?"<>|]', '_', candidate)

    if len(candidate) <= 31:
        return candidate
    else:
        suffix = f"_{clean_sn[:6]}_{clean_crash}"
        prefix_len = 31 - len(suffix)
        return f"{short_file[:prefix_len]}{suffix}"


def extract_time_key(filename):
    """提取时间戳"""
    match = re.search(r"\[(\d{4}-\d{2}-\d{2}\s+[\d\-]+)\]", filename)
    return match.group(1) if match else filename


def add_line_chart_to_excel_sheet(ws, sheet_title, max_cols, max_rows):
    """为 Excel Sheet 绘制折线图 (高 15cm, 宽 30cm)"""
    if max_cols < 1 or max_rows < 2:
        return

    chart = LineChart()
    chart.title = f"{sheet_title} - 波形对比图"
    chart.style = 10
    chart.y_axis.title = "加速度 (g)"
    chart.x_axis.title = "采样点序号 (Points)"

    data = Reference(ws, min_col=1, min_row=1, max_col=max_cols, max_row=max_rows)
    chart.add_data(data, titles_from_data=True)

    # 图表尺寸设置 (高 15cm, 宽 30cm)
    chart.height = 15.00
    chart.width = 30.00

    chart_col_letter = get_column_letter(max_cols + 2)
    ws.add_chart(chart, f"{chart_col_letter}2")


def write_columns_and_chart_to_sheet(ws, sheet_title, file_list, target_dir, header_font, header_fill, header_align):
    """向 Sheet 写入列数据并插图"""
    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 40
    col_idx = 1
    max_data_rows = 1

    for jf in file_list:
        json_path = os.path.join(target_dir, jf)
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                content = json.load(f)

            array_vars = [(k, v) for k, v in content.items() if isinstance(v, list) and len(v) > 0]
            if not array_vars:
                continue

            for var_name, arr_values in array_vars:
                header_title = jf if len(array_vars) == 1 else f"{jf}\n({var_name})"

                cell_header = ws.cell(row=1, column=col_idx, value=header_title)
                cell_header.font = header_font
                cell_header.fill = header_fill
                cell_header.alignment = header_align

                for row_offset, val in enumerate(arr_values, start=2):
                    clean_val = val if val is not None else ""
                    ws.cell(row=row_offset, column=col_idx, value=clean_val)

                if (len(arr_values) + 1) > max_data_rows:
                    max_data_rows = len(arr_values) + 1

                col_letter = get_column_letter(col_idx)
                ws.column_dimensions[col_letter].width = 28
                col_idx += 1
        except Exception as e:
            print(f"  │  └─ 读取 JSON 失败 [{jf}]: {str(e)}")

    written_cols = col_idx - 1
    if written_cols >= 1 and max_data_rows >= 2:
        add_line_chart_to_excel_sheet(ws, sheet_title, max_cols=written_cols, max_rows=max_data_rows)

    return written_cols


def process_json_folder_to_excel(target_dir):
    """将指定目录下的 JSON 导出为 Excel 文件"""
    json_files = [f for f in os.listdir(target_dir) if f.lower().endswith('.json')]
    if not json_files:
        return None

    json_files.sort(key=extract_time_key)
    folder_name = os.path.basename(os.path.normpath(target_dir))
    output_excel_path = os.path.join(target_dir, f"{folder_name}_waveform_vertical.xlsx")

    print(f"  ├─ [2/3 Excel生成] 找到 {len(json_files)} 个 JSON，开始构建 Excel...")

    earliest_x_per_ip = {}
    earliest_y_per_ip = {}
    sheets_files_map = {}

    for jf in json_files:
        base_file, time_str, acu_sn, crash_type = parse_meta_from_json_filename(jf)
        if "Crash X" in crash_type and base_file not in earliest_x_per_ip:
            earliest_x_per_ip[base_file] = jf
        elif "Crash Y" in crash_type and base_file not in earliest_y_per_ip:
            earliest_y_per_ip[base_file] = jf

        sheet_title = make_smart_sheet_title(base_file, acu_sn, crash_type)
        sheets_files_map.setdefault(sheet_title, []).append(jf)

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_font = Font(name="微软雅黑", size=9, bold=True, color="1F497D")
    header_fill = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    summary_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")

    # 1. 汇总 Sheet X
    sheet_name_summary_x = f"{folder_name}_各IP第一条_X"[:31]
    ws_summary_x = wb.create_sheet(title=sheet_name_summary_x)
    first_x_files = [earliest_x_per_ip[ip] for ip in sorted(earliest_x_per_ip.keys())]
    write_columns_and_chart_to_sheet(ws_summary_x, sheet_name_summary_x, first_x_files, target_dir, header_font,
                                     summary_fill, header_align)

    # 2. 汇总 Sheet Y
    sheet_name_summary_y = f"{folder_name}_各IP第一条_Y"[:31]
    ws_summary_y = wb.create_sheet(title=sheet_name_summary_y)
    first_y_files = [earliest_y_per_ip[ip] for ip in sorted(earliest_y_per_ip.keys())]
    write_columns_and_chart_to_sheet(ws_summary_y, sheet_name_summary_y, first_y_files, target_dir, header_font,
                                     summary_fill, header_align)

    # 3. 详细 Sheet
    for sheet_name, files_list in sheets_files_map.items():
        ws = wb.create_sheet(title=sheet_name)
        write_columns_and_chart_to_sheet(ws, sheet_name, files_list, target_dir, header_font, header_fill, header_align)

    # 保存
    try:
        wb.save(output_excel_path)
        print(f"  │  └─ Excel 生成成功 -> {output_excel_path}")
    except PermissionError:
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        output_excel_path = os.path.join(target_dir, f"{folder_name}_waveform_vertical_{timestamp}.xlsx")
        wb.save(output_excel_path)
        print(f"  │  └─ 主文件占用，已生成备用 Excel -> {output_excel_path}")

    return output_excel_path


# =============================================================================
# 第三部分：Excel 数据汇总 -> Word 深度报告生成模块
# =============================================================================

def setup_document_landscape(doc):
    """设置 Word 页面：35cm 宽, 25cm 高"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(35)
    section.page_height = Cm(25)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def analyze_waveform_data(df):
    """提取最高点/最低点，计算均值与 +-5% 合格门限"""
    cols = df.columns
    if len(cols) == 0:
        return None

    p1_list, t1_list = [], []
    for col in cols:
        vals = pd.to_numeric(df[col], errors='coerce').dropna().values
        if len(vals) == 0:
            continue
        p1_list.append(float(np.max(vals)))
        t1_list.append(float(np.min(vals)))

    if not p1_list:
        return None

    def check_metric_range(vals_list):
        avg = float(np.mean(vals_list))
        tol = abs(avg) * 0.05  # 正负 5%
        min_b, max_b = avg - tol, avg + tol
        abnormal = [i + 1 for i, v in enumerate(vals_list) if not (min_b <= v <= max_b)]
        return avg, min_b, max_b, (len(abnormal) == 0), abnormal

    p1_avg, p1_min, p1_max, p1_ok, p1_ab = check_metric_range(p1_list)
    t1_avg, t1_min, t1_max, t1_ok, t1_ab = check_metric_range(t1_list)

    return {
        "p1": {"list": p1_list, "avg": p1_avg, "min": p1_min, "max": p1_max, "ok": p1_ok, "abnormal": p1_ab},
        "t1": {"list": t1_list, "avg": t1_avg, "min": t1_min, "max": t1_max, "ok": t1_ok, "abnormal": t1_ab},
        "is_normal": (p1_ok and t1_ok),
        "series_count": len(p1_list)
    }


def generate_chart_image_for_word(df, title, temp_dir):
    """绘制高 15cm, 宽 30cm 的折线图图片供 Word 使用"""
    fig_width = 30 / 2.54
    fig_height = 15 / 2.54

    fig, ax = plt.subplots(figsize=(fig_width, fig_height), dpi=150)

    for col in df.columns:
        series_data = pd.to_numeric(df[col], errors='coerce').dropna()
        if len(series_data) > 0:
            ax.plot(series_data.index + 1, series_data.values, label=str(col)[:40], linewidth=1.5)

    ax.set_title(f"{title} - 脉冲响应对比图", fontsize=15, fontweight='bold', pad=12)
    ax.set_xlabel("采样点序号 (Points)", fontsize=11)
    ax.set_ylabel("加速度物理幅值 (g)", fontsize=11)
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.axhline(0, color='black', linewidth=0.8, linestyle='-')

    if len(df.columns) <= 10:
        ax.legend(loc='upper right', fontsize=8.5, frameon=True)
    else:
        ax.legend(loc='upper right', bbox_to_anchor=(1.12, 1.0), fontsize=7.5, frameon=True)

    plt.tight_layout()

    safe_title = re.sub(r'[\/:*?"<>|]', '_', title)
    img_path = os.path.join(temp_dir, f"chart_{safe_title}.png")
    plt.savefig(img_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return img_path


def add_summary_table_to_word(doc, stats):
    """插入精确总宽度 30cm 的居中对比分析表格"""
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    col_widths = [Cm(6.0), Cm(5.0), Cm(7.5), Cm(6.5), Cm(5.0)]

    headers = ["极值指标项", "实测均值 (g)", "正负 5% 合格区间范围", "超差通道", "单项评估"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        shading = parse_xml(r'<w:shd {} w:fill="1F497D"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

    def get_status_str(item):
        return "合格 (正常)" if item['ok'] else "超差 (异常)"

    def get_ab_str(item):
        return "无" if item['ok'] else f"通道 {item['abnormal']}"

    data_rows = [
        ("最高点 (峰值)", f"{stats['p1']['avg']:.2f}", f"{stats['p1']['min']:.3f} ～ {stats['p1']['max']:.3f}",
         get_ab_str(stats['p1']), get_status_str(stats['p1'])),
        ("最低点 (谷值)", f"{stats['t1']['avg']:.2f}", f"{stats['t1']['min']:.3f} ～ {stats['t1']['max']:.3f}",
         get_ab_str(stats['t1']), get_status_str(stats['t1'])),
        ("通道数量", f"{stats['series_count']} 条", "综合判定结论", "全在区间内" if stats['is_normal'] else "存在超差",
         "【正常】" if stats['is_normal'] else "【异常】")
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9.0)

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph()


def extract_data_num(path):
    """自然数字排序函数"""
    match = re.search(r'data(\d+)', path, re.IGNORECASE)
    return int(match.group(1)) if match else 999


# =============================================================================
# 第四部分：主控制流水线 (全流程一键调度)
# =============================================================================

def run_full_pipeline(xy_root_dir, output_docx_path):
    if not os.path.exists(xy_root_dir):
        print(f"错误: 根目录不存在 -> {xy_root_dir}")
        return

    print("=" * 60)
    print("【开始全流程一键处理】扫描目录:", xy_root_dir)
    print("=" * 60)

    subdirs = [os.path.join(xy_root_dir, d) for d in os.listdir(xy_root_dir)
               if os.path.isdir(os.path.join(xy_root_dir, d))]
    subdirs.sort(key=extract_data_num)

    excel_files_generated = []

    for sdir in subdirs:
        folder_name = os.path.basename(os.path.normpath(sdir))
        print(f"\n[处理分组目录]: {folder_name}")

        process_txt_folder(sdir)

        excel_file = process_json_folder_to_excel(sdir)
        if excel_file:
            excel_files_generated.append(excel_file)

    if not excel_files_generated:
        print("\n未找到或生成任何 Excel 文件，终止生成 Word 报告。")
        return

    print("\n" + "=" * 60)
    print(f"  ├─ [3/3 Word报告生成] 开始汇总 {len(excel_files_generated)} 个 Excel 数据包...")

    doc = Document()
    setup_document_landscape(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("气囊控制器 (ACM/ACU) 冲击测试波形全角度综合汇总分析报告")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    desc_p = doc.add_paragraph()
    desc_p.add_run("试验环境与判定规则说明：\n").bold = True
    desc_p.add_run("1. 涵盖多角度（如 15°、30°、45°、60°、75°、90° 等）全套冲击测试数据。\n")
    desc_p.add_run(
        "2. 槽位与 IP 映射：1号槽位(172.16.3.171)、2号槽位(172.16.3.172)、3号槽位(172.16.3.173)、4号槽位(172.16.3.174)。\n")
    desc_p.add_run(
        "3. 判定规则：分别提取最高点与最低点计算均值及 [均值×95% ~ 平均值×105%] 合格区间，若所有测试值全落在各自区间内判定为【正常】，否则判定为【异常】。\n")
    doc.add_paragraph()

    temp_dir = tempfile.mkdtemp()

    for excel_path in excel_files_generated:
        rel_folder = os.path.basename(os.path.dirname(excel_path))
        excel_file = os.path.basename(excel_path)

        h1 = doc.add_heading(f"测试组 [{rel_folder}] : {excel_file}", level=1)
        h1.runs[0].font.color.rgb = RGBColor(31, 73, 125)
        h1.runs[0].font.size = Pt(16)

        try:
            xl = pd.ExcelFile(excel_path)
        except Exception as e:
            print(f"  │  └─ 读取 Excel 失败 [{excel_file}]: {str(e)}")
            continue

        for sheet_name in xl.sheet_names:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            if df.empty or len(df.columns) == 0:
                continue

            h2 = doc.add_heading(f"波形对比分组: {sheet_name}", level=2)
            h2.runs[0].font.color.rgb = RGBColor(68, 114, 196)
            h2.runs[0].font.size = Pt(13)

            stats = analyze_waveform_data(df)
            if not stats:
                continue

            chart_img_path = generate_chart_image_for_word(df, f"{rel_folder}_{sheet_name}", temp_dir)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(chart_img_path, height=Cm(15), width=Cm(30))

            p1_str = ", ".join([f"{v:.1f}" for v in stats['p1']['list']])
            t1_str = ", ".join([f"{v:.1f}" for v in stats['t1']['list']])

            status_tag = "【正常】" if stats['is_normal'] else "【异常】"
            eval_desc = "最高点与最低点实测值均 100% 落在各自 ±5% 合格区间内，波动平稳，一致性良好。" if stats[
                'is_normal'] else "部分指标测试值超出 ±5% 门限区间，存在偏差，请关注对应超差通道。"

            narrative = (
                f"【最高点与最低点正负 5% 区间判定】\n"
                f"本组（{sheet_name}）共包含 {stats['series_count']} 条对比波形。针对最高点与最低点的判定如下：\n"
                f"1. 最高点 (峰值)：测试值 [{p1_str}]，计算均值为 {stats['p1']['avg']:.2f} g，正负 5% 合格区间为 [{stats['p1']['min']:.3f} ～ {stats['p1']['max']:.3f}] g → 单项判定: {'【合格】' if stats['p1']['ok'] else '【超差】'}；\n"
                f"2. 最低点 (谷值)：测试值 [{t1_str}]，计算均值为 {stats['t1']['avg']:.2f} g，正负 5% 合格区间为 [{stats['t1']['min']:.3f} ～ {stats['t1']['max']:.3f}] g → 单项判定: {'【合格】' if stats['t1']['ok'] else '【超差】'}。\n"
                f"综合判定结论：{status_tag}。{eval_desc}"
            )

            analysis_p = doc.add_paragraph()
            analysis_p.paragraph_format.line_spacing = 1.25
            run_narrative = analysis_p.add_run(narrative)
            run_narrative.font.size = Pt(10.0)

            add_summary_table_to_word(doc, stats)

    try:
        doc.save(output_docx_path)
        print("\n" + "=" * 60)
        print(f"🎉 全部处理成功！综合 Word 分析报告已生成至 Crash_XY 根目录：\n{output_docx_path}")
        print("=" * 60)
    except PermissionError:
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        backup_docx = os.path.join(xy_root_dir, f"冲击测试波形全角度汇总分析报告_{timestamp}.docx")
        print(f"\n[⚠️ 警告] 目标 Word 文件正被 WPS 打开，已存至备用文件：\n{backup_docx}")
        doc.save(backup_docx)


if __name__ == "__main__":
    # 【相对路径】：自动将当前脚本所在的目录定为 Crash_XY 根目录
    XY_ROOT_DIR = os.path.dirname(os.path.abspath(__file__))

    OUTPUT_DOCX = os.path.join(XY_ROOT_DIR, "冲击测试波形全角度汇总分析报告.docx")

    run_full_pipeline(XY_ROOT_DIR, OUTPUT_DOCX)